#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    set_table_borders(table)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        run.font.size = Pt(12)


def process_docx(docx_path, output_path=None):
    doc = Document(docx_path)
    if not doc.tables:
        target = output_path or docx_path
        if output_path and output_path != docx_path:
            doc.save(target)
            print(f"文档中无表格，已直接输出未改写内容: {target}")
        else:
            print(f"文档中无表格，跳过后处理: {docx_path}")
        return

    for table in doc.tables:
        format_table(table)
    target = output_path or docx_path
    doc.save(target)
    print(f"已处理WORD表格边框: {target}")


def main():
    if len(sys.argv) < 2:
        print("用法: python postprocess_docx.py <input.docx> [output.docx]")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    process_docx(sys.argv[1], output_path)


if __name__ == "__main__":
    main()
